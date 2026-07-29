import hashlib
import io
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as WordDocument
from fastapi import HTTPException, UploadFile, status
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
SUPPORTED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
}
HEADING_PATTERN = re.compile(r"^[A-Z][A-Z0-9 &/+\\-]{2,60}:?$")


@dataclass(frozen=True)
class PreparedChunk:
    sequence: int
    section: str | None
    content: str
    start_character: int
    end_character: int
    token_estimate: int


@dataclass(frozen=True)
class PreparedDocument:
    filename: str
    content_type: str
    sha256: str
    text: str
    metadata: dict
    chunks: list[PreparedChunk]


def _normalise_text(text: str) -> str:
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lines)).strip()


def _extract_text(content: bytes, extension: str) -> tuple[str, dict]:
    try:
        if extension == ".pdf":
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages), {
                "page_count": len(reader.pages)
            }
        if extension == ".docx":
            document = WordDocument(io.BytesIO(content))
            return "\n".join(paragraph.text for paragraph in document.paragraphs), {
                "paragraph_count": len(document.paragraphs)
            }
        return content.decode("utf-8"), {}
    except (UnicodeDecodeError, ValueError, OSError) as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="The uploaded document could not be parsed.",
        ) from exc


def _section_at(text: str, position: int) -> str | None:
    section = None
    for line in text[:position].splitlines():
        if HEADING_PATTERN.fullmatch(line.strip()):
            section = line.strip().rstrip(":").title()
    return section


def build_chunks(text: str, chunk_size: int, overlap: int) -> list[PreparedChunk]:
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            boundary = text.rfind("\n\n", start, end)
            if boundary > start + chunk_size // 2:
                end = boundary
        content = text[start:end].strip()
        if content:
            actual_start = text.find(content, start, end)
            chunks.append(
                PreparedChunk(
                    sequence=len(chunks),
                    section=_section_at(text, actual_start),
                    content=content,
                    start_character=actual_start,
                    end_character=actual_start + len(content),
                    token_estimate=max(1, (len(content) + 3) // 4),
                )
            )
        if end >= len(text):
            break
        start = max(start + 1, end - overlap)
    return chunks


async def prepare_document(
    upload: UploadFile, *, max_bytes: int, chunk_size: int, chunk_overlap: int
) -> PreparedDocument:
    filename = Path(upload.filename or "document").name
    extension = Path(filename).suffix.lower()
    content_type = (upload.content_type or "").lower()
    if extension not in SUPPORTED_EXTENSIONS or content_type not in SUPPORTED_CONTENT_TYPES:
        raise HTTPException(
            status_code=415, detail="Only PDF, DOCX, and UTF-8 TXT documents are supported."
        )
    content = await upload.read(max_bytes + 1)
    await upload.close()
    if not content:
        raise HTTPException(status_code=422, detail="The uploaded document is empty.")
    if len(content) > max_bytes:
        raise HTTPException(status_code=413, detail="The uploaded document is too large.")
    raw_text, metadata = _extract_text(content, extension)
    text = _normalise_text(raw_text)
    if not text:
        raise HTTPException(status_code=422, detail="No readable text was found.")
    metadata.update({"extension": extension, "byte_size": len(content)})
    return PreparedDocument(
        filename=filename,
        content_type=content_type,
        sha256=hashlib.sha256(content).hexdigest(),
        text=text,
        metadata=metadata,
        chunks=build_chunks(text, chunk_size, chunk_overlap),
    )
