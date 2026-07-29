from io import BytesIO

from docx import Document as WordDocument
from fastapi.testclient import TestClient

from app.services.document_ingestion import build_chunks


def test_upload_and_retrieve_text_document(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents",
        data={"document_type": "resume"},
        files={
            "file": (
                "resume.txt",
                b"EXPERIENCE\nBuilt LangGraph interview agents with FastAPI.",
                "text/plain",
            )
        },
    )
    assert response.status_code == 201
    payload = response.json()
    assert payload["document_type"] == "resume"
    assert payload["chunk_count"] == 1

    detail = client.get(f"/api/v1/documents/{payload['id']}")
    assert detail.status_code == 200
    assert detail.json()["chunks"][0]["section"] == "Experience"
    assert "LangGraph" in detail.json()["chunks"][0]["content"]


def test_upload_docx(client: TestClient) -> None:
    document = WordDocument()
    document.add_heading("Requirements", level=1)
    document.add_paragraph("Build production RAG pipelines using Python and FastAPI.")
    content = BytesIO()
    document.save(content)
    response = client.post(
        "/api/v1/documents",
        data={"document_type": "job_description"},
        files={
            "file": (
                "role.docx",
                content.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    assert response.json()["document_metadata"]["paragraph_count"] == 2


def test_rejects_unsupported_file(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents",
        data={"document_type": "resume"},
        files={"file": ("resume.exe", b"not a document", "application/octet-stream")},
    )
    assert response.status_code == 415


def test_rejects_missing_interview_session(client: TestClient) -> None:
    response = client.post(
        "/api/v1/documents",
        data={
            "document_type": "resume",
            "interview_session_id": "00000000-0000-0000-0000-000000000000",
        },
        files={"file": ("resume.txt", b"Python developer", "text/plain")},
    )
    assert response.status_code == 404


def test_chunking_is_deterministic_and_overlapping() -> None:
    text = "EXPERIENCE\n" + ("Built reliable AI systems. " * 30)
    chunks = build_chunks(text, chunk_size=180, overlap=30)
    assert len(chunks) > 1
    assert [chunk.sequence for chunk in chunks] == list(range(len(chunks)))
    assert chunks[1].start_character < chunks[0].end_character
